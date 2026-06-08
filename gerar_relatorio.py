from pathlib import Path
from itertools import product

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Relatorio_Mission_Control_AI.docx"
DIAGRAM = ROOT / "circuito_logico.png"

NAVY = "102A43"
BLUE = "176B87"
CYAN = "35B7C9"
LIGHT = "EAF4F7"
GRAY = "52667A"
RED = "B4232C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def style_table(table, header=True, font_size=9):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                for run in p.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(font_size)
            if header and r_idx == 0:
                set_cell_shading(cell, NAVY)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
        if header and r_idx == 0:
            set_repeat_table_header(row)


def draw_circuit(path):
    width, height = 1500, 750
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    try:
        regular = ImageFont.truetype("arial.ttf", 25)
        bold = ImageFont.truetype("arialbd.ttf", 27)
        small = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        regular = bold = small = ImageFont.load_default()

    def wire(points, color="#173B57", size=5):
        d.line(points, fill=color, width=size, joint="curve")

    def gate(x, y, label, formula, color="#176B87"):
        d.rounded_rectangle((x, y, x + 190, y + 88), radius=18, fill="#EAF4F7", outline=color, width=4)
        d.text((x + 95, y + 25), label, font=bold, fill="#102A43", anchor="mm")
        d.text((x + 95, y + 62), formula, font=small, fill="#52667A", anchor="mm")
        return [(x, y + 29), (x, y + 59)], (x + 190, y + 44)

    def terminal(x, y, key):
        d.ellipse((x - 16, y - 16, x + 16, y + 16), fill="#35B7C9", outline="#102A43", width=2)
        d.text((x, y), key, font=bold, fill="white", anchor="mm")

    d.text((55, 35), "Circuito logico simplificado: X = AC + EF + D(B + ¬A)", font=bold, fill="#102A43")
    and_ac_in, and_ac_out = gate(280, 85, "AND", "AC")
    and_ef_in, and_ef_out = gate(280, 225, "AND", "EF")
    inv_in, inv_out = gate(280, 365, "NOT", "¬A")
    or_ba_in, or_ba_out = gate(560, 435, "OR", "B + ¬A")
    and_d_in, and_d_out = gate(835, 505, "AND", "D(B + ¬A)")
    or_terms_in, or_terms_out = gate(835, 155, "OR", "AC + EF")
    or_final_in, or_final_out = gate(1110, 325, "OR", "X", "#B4232C")

    for key, x, y in [
        ("A", 70, 114), ("C", 70, 144),
        ("E", 70, 254), ("F", 70, 284),
        ("A", 70, 394), ("B", 350, 464), ("D", 625, 564),
    ]:
        terminal(x, y, key)

    wire([(86, 114), and_ac_in[0]])
    wire([(86, 144), and_ac_in[1]])
    wire([(86, 254), and_ef_in[0]])
    wire([(86, 284), and_ef_in[1]])
    wire([(86, 394), inv_in[0]])
    wire([inv_out, (510, 409), (510, 494), or_ba_in[1]])
    wire([(366, 464), or_ba_in[0]])
    wire([or_ba_out, (790, 479), (790, 534), and_d_in[0]])
    wire([(641, 564), and_d_in[1]])
    wire([and_ac_out, (760, 129), (760, 184), or_terms_in[0]])
    wire([and_ef_out, (720, 269), (720, 214), or_terms_in[1]])
    wire([or_terms_out, (1065, 199), (1065, 354), or_final_in[0]])
    wire([and_d_out, (1065, 549), (1065, 384), or_final_in[1]])
    wire([or_final_out, (1370, 369)])
    d.ellipse((1370, 329, 1450, 409), fill="#FF5364", outline="#8D1721", width=5)
    d.text((1410, 369), "X", font=bold, fill="white", anchor="mm")
    d.text((1360, 430), "LED DE ALERTA", font=small, fill="#B4232C", anchor="mm")
    d.text((55, 700), "Implementacao: U1 = 74HC04 (NOT), U2 = 74HC32 (OR), U3 = 74HC08 (AND). Alimentacao: +5 V.", font=regular, fill="#52667A")
    img.save(path, quality=95)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 107, 135)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def evaluate(a, b, c, d, e, f):
    return int((a and c) or (e and f) or (d and (b or not a)))


def build_document():
    draw_circuit(DIAGRAM)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(16, 42, 67)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 6),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11, GRAY, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    styles["List Bullet"].font.name = "Aptos"
    styles["List Bullet"].font.size = Pt(10.5)
    styles["List Bullet"].paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.text = "MISSION CONTROL AI  |  DOCUMENTACAO TECNICA"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Ciencia da Computacao · Turma 1CCPF")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("MISSION CONTROL AI")
    r.font.name = "Aptos Display"
    r.font.size = Pt(29)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Sistema de Alerta por Logica Digital para Missao Espacial Experimental")
    r.font.name = "Aptos Display"
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    fields = [
        ("Aluno(s)", "____________________________________________"),
        ("Turma", "1CCPF"),
        ("Professor", "Mauricio Neto"),
        ("Data", "08/06/2026"),
    ]
    for row, (label, value) in zip(meta.rows, fields):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT)
        row.cells[0].paragraphs[0].runs[0].bold = True
    set_table_geometry(meta, [1800, 7560])
    style_table(meta, header=False, font_size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("OBJETIVO DO PROJETO")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(CYAN)
    p = doc.add_paragraph(
        "Monitorar seis condicoes digitais da nave e acionar a saida X = 1 quando uma combinacao representar risco operacional. "
        "O sistema integra expressao booleana, simplificacao, tabela verdade completa, circuito com CIs e simulacao interativa."
    )
    p.paragraph_format.space_after = Pt(0)

    doc.add_page_break()
    add_heading(doc, "1. Contexto operacional", 1)
    doc.add_paragraph(
        "Durante uma missao espacial experimental, sensores e subsistemas enviam estados binarios ao modulo de decisao. "
        "O valor 0 indica funcionamento normal; o valor 1 indica falha, risco ou alerta. A saida X alimenta um LED vermelho "
        "e pode ser conectada a uma sirene ou ao software da central de controle."
    )

    add_heading(doc, "2. Variaveis de entrada", 1)
    variables = [
        ("A", "Falha de comunicacao", "1 quando o enlace com a nave esta indisponivel."),
        ("B", "Temperatura interna critica", "1 quando a temperatura sai da faixa segura."),
        ("C", "Baixo nivel de energia", "1 quando a reserva energetica fica abaixo do limite."),
        ("D", "Falha em modulo operacional", "1 quando um modulo essencial apresenta defeito."),
        ("E", "Perda de estabilidade", "1 quando a orientacao ou atitude da nave e comprometida."),
        ("F", "Alerta da inteligencia artificial", "1 quando o monitor inteligente detecta um padrao de risco."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Variavel"
    table.rows[0].cells[1].text = "Condicao monitorada"
    table.rows[0].cells[2].text = "Significado do nivel logico 1"
    for item in variables:
        cells = table.add_row().cells
        for cell, value in zip(cells, item):
            cell.text = value
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, [1050, 3000, 5310])
    style_table(table, font_size=9.5)

    add_heading(doc, "3. Regras de acionamento", 1)
    add_bullet(doc, "R1: falha de comunicacao E baixo nivel de energia: A · C.")
    add_bullet(doc, "R2: temperatura critica E falha em modulo: B · D.")
    add_bullet(doc, "R3: alerta da IA E perda de estabilidade: E · F.")
    add_bullet(doc, "R4: falha de modulo quando NAO ha falha de comunicacao: D · ¬A.")
    add_bullet(doc, "R5: condicao redundante de alta criticidade: A · C · D.")

    add_heading(doc, "4. Expressao booleana completa", 1)
    add_formula(doc, "X = (A · C) + (B · D) + (E · F) + (D · ¬A) + (A · C · D)")
    doc.add_paragraph(
        "Os simbolos ·, + e ¬ representam, respectivamente, as portas AND, OR e NOT. "
        "Assim, as tres operacoes obrigatorias aparecem explicitamente no projeto."
    )

    add_heading(doc, "5. Simplificacao", 1)
    doc.add_paragraph("Aplicando a lei da absorcao, AC + ACD = AC:")
    add_formula(doc, "X = AC + BD + EF + D¬A")
    doc.add_paragraph("Fatorando os termos que contem D:")
    add_formula(doc, "X = AC + EF + D(B + ¬A)")
    doc.add_paragraph(
        "A ultima forma e usada no circuito por exigir menos portas. A simplificacao nao altera nenhuma linha da tabela verdade."
    )

    doc.add_page_break()
    add_heading(doc, "6. Tabela verdade completa", 1)
    doc.add_paragraph(
        "As 64 combinacoes possiveis para seis entradas estao listadas abaixo. A coluna Estado destaca as situacoes em que o alerta e acionado."
    )
    truth = doc.add_table(rows=1, cols=9)
    truth.style = "Table Grid"
    headers = ["#", "A", "B", "C", "D", "E", "F", "X", "Estado"]
    for cell, text in zip(truth.rows[0].cells, headers):
        cell.text = text
    for idx, values in enumerate(product([0, 1], repeat=6), start=1):
        x = evaluate(*values)
        row = truth.add_row()
        data = [idx, *values, x, "ALERTA" if x else "NORMAL"]
        for cell, value in zip(row.cells, data):
            cell.text = str(value)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if x:
            set_cell_shading(row.cells[-1], "FCE8EA")
            row.cells[-1].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(RED)
            row.cells[-1].paragraphs[0].runs[0].bold = True
    set_table_geometry(truth, [530, 650, 650, 650, 650, 650, 650, 650, 3280])
    style_table(truth, font_size=8)

    doc.add_page_break()
    add_heading(doc, "7. Circuito logico digital", 1)
    doc.add_paragraph(
        "O circuito implementa diretamente X = AC + EF + D(B + ¬A). Os sinais intermediarios sao combinados por tres CIs da familia 74HC."
    )
    doc.add_picture(str(DIAGRAM), width=Inches(6.65))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "7.1 Lista de componentes", 2)
    components = [
        ("U1", "74HC04", "6 inversores NOT; um deles gera ¬A."),
        ("U2", "74HC32", "4 portas OR de duas entradas; gera B + ¬A e a soma final."),
        ("U3", "74HC08", "4 portas AND de duas entradas; gera AC, EF e D(B + ¬A)."),
        ("LED1", "LED vermelho", "Indicador visual da saida X."),
        ("R1", "330 ohms", "Limita a corrente do LED."),
        ("R2-R7", "10 kohms", "Resistores pull-down das seis entradas."),
        ("S1-S6", "Chaves", "Selecionam 0 ou 1 para A, B, C, D, E e F."),
        ("-", "Protoboard e fonte de 5 V", "Montagem, alimentacao e fios de interligacao."),
    ]
    comp = doc.add_table(rows=1, cols=3)
    comp.style = "Table Grid"
    for cell, text in zip(comp.rows[0].cells, ["Ref.", "Componente", "Funcao"]):
        cell.text = text
    for item in components:
        cells = comp.add_row().cells
        for cell, value in zip(cells, item):
            cell.text = value
    set_table_geometry(comp, [950, 2300, 6110])
    style_table(comp, font_size=9)

    add_heading(doc, "7.2 Ligacoes dos CIs", 2)
    add_bullet(doc, "Todos os CIs: pino 14 em +5 V e pino 7 em GND. Usar um capacitor de 100 nF entre esses pinos em cada CI.")
    add_bullet(doc, "U1 74HC04: A no pino 1; saida ¬A no pino 2.")
    add_bullet(doc, "U2 74HC32: B nos pinos 1 e ¬A no 2, produzindo B + ¬A no pino 3.")
    add_bullet(doc, "U3 74HC08: A/C nos pinos 1/2 (AC no 3); E/F nos pinos 4/5 (EF no 6); D e U2-3 nos pinos 9/10 (termo D(B + ¬A) no 8).")
    add_bullet(doc, "U2 74HC32: AC e EF nos pinos 4/5 (soma no 6); U2-6 e U3-8 nos pinos 9/10; saida X no pino 8.")
    add_bullet(doc, "Ligar U2-8 ao anodo do LED; ligar o catodo ao resistor de 330 ohms e depois ao GND.")
    add_bullet(doc, "Entradas de portas nao utilizadas nunca devem ficar flutuando: conecta-las ao GND.")

    add_heading(doc, "8. Funcionamento do alerta", 1)
    doc.add_paragraph(
        "O alerta X = 1 e acionado por qualquer um dos tres caminhos da expressao simplificada: "
        "(1) comunicacao falha junto com energia baixa; (2) a IA detecta risco enquanto a nave perde estabilidade; "
        "ou (3) um modulo falha e, simultaneamente, ocorre temperatura critica ou a comunicacao permanece funcional. "
        "Quando nenhum caminho e verdadeiro, X = 0 e o LED fica apagado."
    )

    add_heading(doc, "9. Procedimento de demonstracao", 1)
    for text in [
        "Abrir o arquivo simulacao_mission_control.html em um navegador.",
        "Testar A=1 e C=1; o alerta deve acender independentemente das demais entradas.",
        "Testar E=1 e F=1; o alerta deve acender.",
        "Testar D=1 e A=0; o alerta deve acender porque ¬A=1.",
        "Testar todas as entradas em 0; o alerta deve permanecer apagado.",
        "Comparar cada teste com a linha destacada da tabela verdade na simulacao.",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(text)

    add_heading(doc, "10. Conclusao", 1)
    doc.add_paragraph(
        "O Mission Control AI atende aos requisitos de monitoramento digital da missao com seis variaveis, expressao completa, "
        "uso de AND/OR/NOT, simplificacao, tabela verdade de 64 linhas, circuito com CIs e demonstracao interativa. "
        "A arquitetura e modular e pode ser ampliada com novos sensores ou integrada a um sistema operacional de controle."
    )

    core = doc.core_properties
    core.title = "Mission Control AI"
    core.subject = "Sistema de alerta por logica digital"
    core.author = "Turma 1CCPF"
    core.keywords = "logica digital, tabela verdade, circuito logico, missao espacial"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
